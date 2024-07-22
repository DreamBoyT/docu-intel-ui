import streamlit as st
from langchain_openai import AzureOpenAIEmbeddings, AzureChatOpenAI
import tempfile
import time
import concurrent.futures
from langchain import LLMChain, PromptTemplate
from langchain.text_splitter import CharacterTextSplitter
from langchain.document_loaders import PyPDFLoader
from langchain.vectorstores import FAISS
from langchain.chains.question_answering import load_qa_chain
from PyPDF2 import PdfReader
from io import BytesIO
from docx import Document as DocxDocument
import base64

# Azure OpenAI API details
azure_api_key = 'c09f91126e51468d88f57cb83a63ee36'
azure_endpoint = 'https://chat-gpt-a1.openai.azure.com/'
azure_api_version = '2024-02-01'
azure_chat_endpoint = 'https://danielingitaraj.openai.azure.com/'
openai_api_key = 'a5c4e09a50dd4e13a69e7ef19d07b48c'

# Initialize Azure OpenAI Embeddings
embed_model = AzureOpenAIEmbeddings(
    model="text-embedding-3-large",
    deployment="text-embedding-3-large",
    api_version="2023-12-01-preview",
    azure_endpoint=azure_endpoint,
    openai_api_key=azure_api_key,
)

# Initialize Azure OpenAI LLM
llm = AzureChatOpenAI(
    openai_api_key=openai_api_key,
    api_version=azure_api_version,
    azure_endpoint=azure_chat_endpoint,
    model="gpt-4",
    base_url=None,
    azure_deployment="GPT4",
    temperature=0.5,  # Adjusted temperature for improved summaries
)

# Text Splitter
text_splitter = CharacterTextSplitter(
    separator="\n",
    chunk_size=1500,
    chunk_overlap=200,
    length_function=len,
)

# Streamlit user interface
st.title("Document Intelligent Application")
pdf_file = st.sidebar.file_uploader("Choose a PDF file", type="pdf")

def extract_text_from_pdf(file):
    reader = PdfReader(file)
    return [(i + 1, page.extract_text()) for i, page in enumerate(reader.pages)]

def create_prompt(page_numbers, combined_text):
    combined_text = combined_text.replace("{", "{{").replace("}", "}}")
    return f"""Generate a concise and informative summary of the main content from these pages, using letters (a), (b), (c), etc. to denote each key point. Allow the summary to unfold naturally, without a predetermined number of points. Focus on capturing the most essential and relevant information, ensuring that each point adds significant value to the summary. Prioritize quality over quantity, and avoid including unnecessary points. Optimize the summary for clarity, coherence, and speed of generation.

    {combined_text}
    """

def summarize_pages(llm, page_numbers, combined_text):
    prompt = create_prompt(page_numbers, combined_text)
    prompt_template = PromptTemplate.from_template(prompt)
    chain = LLMChain(llm=llm, prompt=prompt_template)
    response = chain.run({"combined_text": combined_text})
    start_page, end_page = page_numbers[0], page_numbers[-1]
    return f"*Pages {start_page}-{end_page}:*\n\n{response.strip()}\n"

def group_texts(texts, group_size=3):
    grouped_texts = []
    for i in range(0, len(texts), group_size):
        group = texts[i:i + group_size]
        page_numbers = [num for num, _ in group]
        combined_text = "\n".join([text for _, text in group])
        grouped_texts.append((page_numbers, combined_text))
    return grouped_texts

def extract_summaries_from_pdf(llm, file, group_size):
    texts = extract_text_from_pdf(file)
    grouped_texts = group_texts(texts, group_size)
    summaries = [None] * len(grouped_texts)
    with concurrent.futures.ThreadPoolExecutor() as executor:
        futures = {executor.submit(summarize_pages, llm, page_numbers, combined_text): idx for idx, (page_numbers, combined_text) in enumerate(grouped_texts)}
        for future in concurrent.futures.as_completed(futures):
            idx = futures[future]
            try:
                summaries[idx] = future.result()
            except Exception as e:
                start_page, end_page = grouped_texts[idx][0][0], grouped_texts[idx][0][-1]
                summaries[idx] = f"*Pages {start_page}-{end_page}:*\n\nError summarizing pages {start_page}-{end_page}: {e}\n"
    return "\n".join(summaries)

def generate_word_file(summaries):
    doc = DocxDocument()
    doc.add_heading('Document Summary', 0)
    for summary in summaries.split("\n\n"):
        doc.add_paragraph(summary)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

if "chat_history" not in st.session_state:
    st.session_state.chat_history = []

def display_chat_history():
    st.subheader("Chat History")
    chat_container = st.container()
    with chat_container:
        for chat in st.session_state.chat_history:
            st.write(chat)
    st.write("End of Chat History")

overall_summary = ""
question_summaries = ""

if pdf_file is not None:
    with tempfile.NamedTemporaryFile(delete=False) as tmp_file:
        tmp_file.write(pdf_file.read())
        pdf_path = tmp_file.name
        loader = PyPDFLoader(pdf_path)
        pages = loader.load_and_split()

    summary_option = st.sidebar.radio(
        "Choose Summary Option",
        options=[
            "Generate 1 Page Summary",
            "Generate 3 Page Summary (default)",
            "Generate 5 Page Summary"
        ],
        index=1
    )

    if summary_option == "Generate 1 Page Summary":
        combined_content = ''.join([p.page_content for p in pages])
        texts = text_splitter.split_text(combined_content)
        overall_summary = extract_summaries_from_pdf(llm, pdf_path, group_size=1)
    elif summary_option == "Generate 3 Page Summary (default)":
        combined_content = ''.join([p.page_content for p in pages])
        texts = text_splitter.split_text(combined_content)
        overall_summary = extract_summaries_from_pdf(llm, pdf_path, group_size=3)
    elif summary_option == "Generate 5 Page Summary":
        combined_content = ''.join([p.page_content for p in pages])
        texts = text_splitter.split_text(combined_content)
        overall_summary = extract_summaries_from_pdf(llm, pdf_path, group_size=5)

    if overall_summary:
        with st.sidebar:
            col1, col2 = st.columns([3, 1])
            with col1:
                st.subheader(pdf_file.name)
                st.write(overall_summary)
            with col2:
                word_file = generate_word_file(overall_summary)
                file_name = pdf_file.name.rsplit('.', 1)[0] + "_summary.docx"
                st.markdown(
                    f"""
                    <style>
                    .download-button {{
                        position: relative;
                        display: inline-block;
                        margin-top: 10px;
                    }}
                    .download-button button {{
                        background-color: #1a1c23;
                        color: white;
                        border: 2px solid #262730;
                        border-radius: 8px;
                        padding: 10px 20px;
                        cursor: pointer;
                        font-size: 14px;
                        transition: background-color 0.3s, border-color 0.3s;
                    }}
                    .download-button button:hover {{
                        border-color: #1a1c23;
                    }}
                    .tooltip {{
                        visibility: hidden;
                        width: 160px;
                        background-color: #555;
                        color: #fff;
                        text-align: center;
                        border-radius: 5px;
                        padding: 5px 0;
                        position: absolute;
                        z-index: 1;
                        bottom: 125%;
                        left: 50%;
                        margin-left: -80px;
                        opacity: 0;
                        transition: opacity 0.3s;
                    }}
                    .download-button:hover .tooltip {{
                        visibility: visible;
                        opacity: 1;
                    }}
                    </style>
                    <div class="download-button">
                        <a href="data:application/octet-stream;base64,{base64.b64encode(word_file.getvalue()).decode()}" download="{file_name}">
                            <button>Download Summary</button>
                            <span class="tooltip" style="background-color: #1a1c23;">Download the summary in Word format</span>
                        </a>
                    </div>
                    """,
                    unsafe_allow_html=True
                )

    question = st.text_input("Enter your question")
    if st.button("Submit"):
        if question:
            combined_content = ''.join([p.page_content for p in pages])
            texts = text_splitter.split_text(combined_content)
            document_search = FAISS.from_texts(texts, embed_model)
            docs = document_search.similarity_search(question)
            chain = load_qa_chain(llm, chain_type="stuff")
            summaries = chain.run(
                input_documents=docs,
                question=question
            )
            st.subheader("Question Answering Result")
            st.write(summaries)
            st.session_state.chat_history.append({"question": question, "answer": summaries})
        else:
            st.warning("Please enter a valid question.")
else:
    time.sleep(35)
    st.warning("No PDF file uploaded")

display_chat_history()
